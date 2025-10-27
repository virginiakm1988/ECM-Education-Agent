"""
Emergency Response Knowledge Base Management
"""
import json
import os
from typing import List, Dict, Any, Optional
from datetime import datetime
import logging

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import PyPDF2
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

class EmergencyKnowledgeBase:
    """
    Manages the emergency response knowledge base with comprehensive emergency management content
    """
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        
    def get_comprehensive_emergency_knowledge(self) -> List[Document]:
        """
        Returns comprehensive emergency management knowledge base
        """
        knowledge_items = [
            # Natural Disasters
            {
                "content": """
                EARTHQUAKE RESPONSE PROCEDURES:
                
                IMMEDIATE ACTIONS (First 2 minutes):
                1. DROP to hands and knees
                2. TAKE COVER under a sturdy desk or table
                3. HOLD ON to your shelter and protect your head/neck
                4. Stay away from windows, mirrors, and heavy objects
                5. If outdoors, move away from buildings, trees, and power lines
                
                AFTER SHAKING STOPS:
                1. Check for injuries and provide first aid
                2. Check for hazards (gas leaks, electrical damage, structural damage)
                3. Turn off utilities if damage is suspected
                4. Evacuate if building is damaged
                5. Stay out of damaged buildings
                6. Listen to emergency broadcasts
                7. Be prepared for aftershocks
                
                COMMUNICATION:
                - Use text messages instead of phone calls
                - Monitor emergency radio frequencies
                - Check in with family/colleagues when safe
                """,
                "category": "Natural Disasters",
                "subcategory": "Earthquake"
            },
            {
                "content": """
                FIRE EMERGENCY PROCEDURES:
                
                DISCOVERY OF FIRE:
                1. Sound the alarm immediately
                2. Call 911 or emergency services
                3. Attempt to extinguish ONLY if fire is small and you have proper equipment
                4. If fire cannot be controlled, evacuate immediately
                
                EVACUATION PROCEDURES:
                1. Use nearest safe exit
                2. Feel doors before opening (if hot, use alternate route)
                3. Stay low if smoke is present
                4. Close doors behind you to slow fire spread
                5. Use stairs, never elevators
                6. Proceed to designated assembly area
                7. Report to floor wardens/emergency coordinators
                
                FIRE EXTINGUISHER USE (PASS Method):
                P - Pull the pin
                A - Aim at base of fire
                S - Squeeze the handle
                S - Sweep from side to side
                
                NEVER fight fire if:
                - Fire is larger than you
                - Smoke is filling the room
                - Fire is spreading rapidly
                - Your exit route is threatened
                """,
                "category": "Natural Disasters",
                "subcategory": "Fire"
            },
            {
                "content": """
                SEVERE WEATHER RESPONSE:
                
                TORNADO WARNING:
                1. Move to lowest floor of building
                2. Go to interior room away from windows
                3. Avoid large roof spans (cafeterias, gyms, auditoriums)
                4. Get under sturdy furniture if possible
                5. Protect head and neck with arms
                6. Stay away from windows and glass
                7. Monitor weather radio for updates
                
                SEVERE THUNDERSTORM:
                1. Move indoors immediately
                2. Stay away from windows
                3. Avoid electrical equipment and plumbing
                4. Do not use landline phones
                5. Wait 30 minutes after last thunder before going outside
                
                FLOODING:
                1. Move to higher ground immediately
                2. Avoid walking/driving through flood water
                3. Turn off utilities if instructed
                4. Do not touch electrical equipment if wet
                5. Listen for evacuation orders
                """,
                "category": "Natural Disasters",
                "subcategory": "Severe Weather"
            },
            
            # Medical Emergencies
            {
                "content": """
                MEDICAL EMERGENCY RESPONSE:
                
                INITIAL ASSESSMENT:
                1. Ensure scene safety
                2. Check responsiveness
                3. Call 911 immediately if serious
                4. Check airway, breathing, circulation (ABC)
                5. Control bleeding if present
                6. Treat for shock
                7. Do not move victim unless in immediate danger
                
                CARDIAC ARREST/CPR:
                1. Check for responsiveness and breathing
                2. Call 911 and request AED
                3. Begin chest compressions:
                   - Place heel of hand on center of chest
                   - Push hard and fast at least 2 inches deep
                   - Rate of 100-120 compressions per minute
                   - Allow complete chest recoil
                4. Provide rescue breaths if trained
                5. Continue until EMS arrives or AED becomes available
                
                CHOKING (Conscious Adult):
                1. Ask "Are you choking?"
                2. Give 5 back blows between shoulder blades
                3. Give 5 abdominal thrusts (Heimlich maneuver)
                4. Continue alternating until object is expelled or person becomes unconscious
                5. If unconscious, begin CPR
                
                SEVERE BLEEDING:
                1. Apply direct pressure with clean cloth
                2. Elevate injured area above heart if possible
                3. Apply pressure to pressure points if bleeding continues
                4. Do not remove objects impaled in wounds
                5. Treat for shock
                """,
                "category": "Medical Emergencies",
                "subcategory": "Life-Threatening"
            },
            
            # Security Incidents
            {
                "content": """
                ACTIVE SHOOTER RESPONSE (RUN-HIDE-FIGHT):
                
                RUN:
                1. Have an escape route and plan in mind
                2. Leave belongings behind
                3. Help others escape if possible
                4. Keep hands visible when evacuating
                5. Follow instructions of police officers
                6. Do not attempt to move wounded people
                7. Call 911 when safe
                
                HIDE:
                1. Hide in area out of shooter's view
                2. Block entry to hiding place and lock doors
                3. Silence cell phones and remain quiet
                4. Turn off lights and close blinds
                5. Barricade doors with heavy furniture
                6. Spread out if in group
                7. Remain calm and wait for law enforcement
                
                FIGHT (Last Resort):
                1. Act as a team if possible
                2. Improvise weapons
                3. Be aggressive and committed
                4. Throw items at shooter
                5. Yell loudly
                6. Attack shooter's ability to see and breathe
                
                WHEN LAW ENFORCEMENT ARRIVES:
                1. Keep hands visible and empty
                2. Follow all commands immediately
                3. Avoid quick movements
                4. Do not point or yell
                5. Know that first officers may not help injured
                """,
                "category": "Security Incidents",
                "subcategory": "Active Threat"
            },
            
            # Infrastructure Failures
            {
                "content": """
                POWER OUTAGE RESPONSE:
                
                IMMEDIATE ACTIONS:
                1. Check if outage is localized or widespread
                2. Turn off electrical equipment to prevent damage from power surges
                3. Keep refrigerators and freezers closed
                4. Use flashlights, not candles
                5. Listen to battery-powered radio for updates
                6. Check on neighbors, especially elderly
                
                EXTENDED OUTAGE PROCEDURES:
                1. Conserve phone battery
                2. Use generators outdoors only
                3. Avoid opening refrigerator/freezer
                4. Dress warmly if heating is affected
                5. Stay hydrated
                6. Monitor for signs of carbon monoxide poisoning
                
                ELEVATOR ENTRAPMENT:
                1. Press alarm button or call button
                2. Use emergency phone if available
                3. Call 911 from cell phone
                4. Do not attempt to force doors open
                5. Do not try to climb out
                6. Remain calm and wait for help
                7. If power fails, emergency lighting should activate
                
                HVAC SYSTEM FAILURE:
                1. Report to facilities management
                2. Open windows if safe and weather permits
                3. Move to areas with better ventilation
                4. Stay hydrated
                5. Monitor for heat-related illness
                6. Consider early dismissal if conditions worsen
                """,
                "category": "Infrastructure Failures",
                "subcategory": "Utilities"
            },
            
            # Communication Systems
            {
                "content": """
                CRISIS COMMUNICATION TEMPLATES:
                
                INITIAL ALERT TEMPLATE:
                "EMERGENCY ALERT: [Type of emergency] occurring at [Location] at [Time]. 
                [Immediate action required]. Emergency services have been notified. 
                Updates will follow every [frequency]. For information: [Contact]."
                
                UPDATE TEMPLATE:
                "EMERGENCY UPDATE [#]: [Current status of situation]. 
                [Actions taken]. [Current instructions for personnel]. 
                [Expected next update time]. [Contact for questions]."
                
                ALL-CLEAR TEMPLATE:
                "ALL-CLEAR: The emergency situation at [Location] has been resolved as of [Time]. 
                [Brief summary of resolution]. Normal operations [will resume/have resumed] at [Time]. 
                [Any follow-up actions required]. Thank you for your cooperation."
                
                MEDIA STATEMENT TEMPLATE:
                "At approximately [Time] on [Date], [Organization] experienced [Type of incident] 
                at [Location]. [Brief factual description]. [Actions taken]. 
                [Current status]. [Cooperation with authorities]. 
                [Contact information for media inquiries]."
                
                FAMILY NOTIFICATION TEMPLATE:
                "This message is to inform you that [Organization] has experienced [Type of emergency]. 
                All personnel have been [accounted for/evacuated safely]. 
                [Any injuries or impacts]. [Current status]. 
                [When normal operations will resume]. [Contact for family questions]."
                """,
                "category": "Crisis Communication",
                "subcategory": "Templates"
            },
            
            # Business Continuity
            {
                "content": """
                BUSINESS CONTINUITY ACTIVATION:
                
                DECISION CRITERIA:
                1. Threat to life safety
                2. Significant property damage
                3. Loss of critical systems for >4 hours
                4. Inability to access primary facility
                5. Loss of key personnel
                6. Regulatory/legal requirements
                7. Significant financial impact
                
                ACTIVATION PROCESS:
                1. Assess situation severity
                2. Notify Business Continuity Team
                3. Activate alternate facilities if needed
                4. Implement communication plan
                5. Deploy recovery teams
                6. Monitor and adjust response
                7. Document all actions
                
                CRITICAL FUNCTIONS PRIORITY:
                1. Life safety and security
                2. Emergency communications
                3. Critical business processes
                4. IT systems and data recovery
                5. Customer service continuity
                6. Financial operations
                7. Regulatory compliance
                8. Supply chain management
                
                RECOVERY PHASES:
                Phase 1 (0-24 hours): Life safety, damage assessment, emergency communications
                Phase 2 (1-7 days): Critical function restoration, alternate site activation
                Phase 3 (1-4 weeks): Full operations restoration, lessons learned
                Phase 4 (1+ months): Return to normal operations, plan updates
                """,
                "category": "Business Continuity",
                "subcategory": "Activation"
            },
            
            # Incident Command System
            {
                "content": """
                INCIDENT COMMAND SYSTEM (ICS) ROLES:
                
                INCIDENT COMMANDER (IC):
                - Overall responsibility for incident management
                - Establishes incident objectives and strategy
                - Approves resource requests
                - Authorizes information release
                - Ensures safety of all personnel
                
                OPERATIONS SECTION CHIEF:
                - Manages tactical operations
                - Develops tactical assignments
                - Coordinates with other agencies
                - Manages resources assigned to operations
                
                PLANNING SECTION CHIEF:
                - Collects and evaluates information
                - Prepares Incident Action Plan
                - Maintains resource status
                - Prepares incident documentation
                
                LOGISTICS SECTION CHIEF:
                - Provides support and resources
                - Manages communications
                - Coordinates medical support
                - Handles procurement and cost analysis
                
                FINANCE/ADMINISTRATION SECTION CHIEF:
                - Tracks incident costs
                - Handles procurement contracts
                - Manages personnel time records
                - Processes claims and compensation
                
                SAFETY OFFICER:
                - Monitors safety conditions
                - Develops safety measures
                - Has authority to stop unsafe acts
                - Investigates accidents
                
                PUBLIC INFORMATION OFFICER:
                - Manages media relations
                - Coordinates information release
                - Handles public inquiries
                - Manages social media communications
                """,
                "category": "Incident Management",
                "subcategory": "ICS Structure"
            }
        ]
        
        documents = []
        for item in knowledge_items:
            # Split long content into chunks
            chunks = self.text_splitter.split_text(item["content"])
            
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "category": item["category"],
                        "subcategory": item["subcategory"],
                        "source": "comprehensive_knowledge_base",
                        "chunk_id": f"{item['subcategory']}_{i}",
                        "timestamp": datetime.now().isoformat()
                    }
                )
                documents.append(doc)
        
        return documents
    
    def process_pdf_document(self, file_path: str, category: str = "Custom") -> List[Document]:
        """
        Process PDF document and extract text for knowledge base
        """
        try:
            documents = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                full_text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    full_text += f"\n--- Page {page_num + 1} ---\n{text}"
                
                # Split text into chunks
                chunks = self.text_splitter.split_text(full_text)
                
                for i, chunk in enumerate(chunks):
                    doc = Document(
                        page_content=chunk,
                        metadata={
                            "category": category,
                            "source": os.path.basename(file_path),
                            "file_type": "pdf",
                            "chunk_id": f"pdf_{i}",
                            "timestamp": datetime.now().isoformat()
                        }
                    )
                    documents.append(doc)
            
            logger.info(f"Processed PDF: {file_path}, created {len(documents)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return []
    
    def process_docx_document(self, file_path: str, category: str = "Custom") -> List[Document]:
        """
        Process Word document and extract text for knowledge base
        """
        try:
            documents = []
            
            doc = DocxDocument(file_path)
            
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += cell.text + " "
                    full_text += "\n"
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(full_text)
            
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "category": category,
                        "source": os.path.basename(file_path),
                        "file_type": "docx",
                        "chunk_id": f"docx_{i}",
                        "timestamp": datetime.now().isoformat()
                    }
                )
                documents.append(doc)
            
            logger.info(f"Processed DOCX: {file_path}, created {len(documents)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return []
    
    def create_custom_knowledge(self, content: str, category: str, subcategory: str = "") -> List[Document]:
        """
        Create knowledge documents from custom content
        """
        chunks = self.text_splitter.split_text(content)
        documents = []
        
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "category": category,
                    "subcategory": subcategory,
                    "source": "custom_input",
                    "chunk_id": f"custom_{i}",
                    "timestamp": datetime.now().isoformat()
                }
            )
            documents.append(doc)
        
        return documents
    
    def get_emergency_contacts_template(self) -> Dict[str, Any]:
        """
        Returns template for emergency contacts
        """
        return {
            "internal_contacts": {
                "emergency_coordinator": {"name": "", "phone": "", "email": "", "role": "Primary Emergency Coordinator"},
                "backup_coordinator": {"name": "", "phone": "", "email": "", "role": "Backup Emergency Coordinator"},
                "facilities_manager": {"name": "", "phone": "", "email": "", "role": "Facilities Management"},
                "it_manager": {"name": "", "phone": "", "email": "", "role": "IT Systems"},
                "hr_manager": {"name": "", "phone": "", "email": "", "role": "Human Resources"},
                "security_manager": {"name": "", "phone": "", "email": "", "role": "Security"}
            },
            "external_contacts": {
                "police": {"phone": "911", "non_emergency": "", "contact_person": ""},
                "fire_department": {"phone": "911", "non_emergency": "", "contact_person": ""},
                "ems": {"phone": "911", "non_emergency": "", "contact_person": ""},
                "hospital": {"name": "", "phone": "", "address": ""},
                "utilities": {
                    "electric": {"company": "", "emergency_phone": ""},
                    "gas": {"company": "", "emergency_phone": ""},
                    "water": {"company": "", "emergency_phone": ""},
                    "telecom": {"company": "", "emergency_phone": ""}
                },
                "emergency_management": {"agency": "", "phone": "", "contact_person": ""}
            },
            "media_contacts": {
                "primary_spokesperson": {"name": "", "phone": "", "email": ""},
                "backup_spokesperson": {"name": "", "phone": "", "email": ""},
                "media_relations": {"company": "", "contact": "", "phone": ""}
            }
        }

# Example usage
if __name__ == "__main__":
    kb = EmergencyKnowledgeBase()
    
    # Get comprehensive knowledge
    docs = kb.get_comprehensive_emergency_knowledge()
    print(f"Created {len(docs)} knowledge documents")
    
    # Get emergency contacts template
    contacts = kb.get_emergency_contacts_template()
    print("\nEmergency Contacts Template:")
    print(json.dumps(contacts, indent=2))
